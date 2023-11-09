import unittest
from docx import Document
import interview_transform
import docx2html
import re
import logging
from lxml import etree
from io import StringIO 



TEST_FILE = './test_interview.docx'

class TestInterviewTransform(unittest.TestCase):
    def setUp(self) -> None:
        self.doc = Document(TEST_FILE)
    
    def test_find_first_para(self):
        first = interview_transform.get_start_para(self.doc)
        self.assertEqual(first, 44)
    
    def test_handle_bf(self):
        first = interview_transform.get_start_para(self.doc)
        handled = interview_transform.handle_bf(self.doc.paragraphs[first])
        self.assertEqual(handled[0:18], 'BOLDPerson OneBOLD')
    
    def test_handle_italics(self):
        first = interview_transform.get_start_para(self.doc)
        handled = interview_transform.handle_italics(self.doc.paragraphs[first])
        self.assertEqual(handled[-40:], 'ITALICSThis text is italicized.ITALICS::')
    

class TestMainFunction(unittest.TestCase):
    def setUp(self) -> None:
        self.doc = Document(TEST_FILE)

    def test_html_length(self):
        html = interview_transform.docx2html(self.doc, TEST_FILE)
        self.assertEqual(len(html), 325)
    
    def test_tree(self):
        html = interview_transform.docx2html(self.doc, TEST_FILE)
        parser = etree.XMLParser(remove_blank_text=True)
        tree   = etree.parse(StringIO(html), parser)
        root = tree.getroot()
        dts = []
        for element in root.iter("*"):
            if element.tag == 'dt':
                dts.append(element.text)
        self.assertEqual(dts[0], 'Person One')
        

unittest.main()